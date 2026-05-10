# -*- coding: utf-8 -*-
"""Generate a detailed Chinese user manual DOCX from an existing DOCX with screenshots.

Usage (PowerShell 7 / CMD both ok):
  python e:\项目\软著文档\generate_user_manual_docx.py \
    --src  e:\项目\软著文档\用户手册.docx \
    --out  e:\项目\软著文档\用户手册_生成版.docx \
    --root e:\项目

It will also write:
  e:\项目\软著文档\用户手册_生成版_截图索引.csv
  e:\项目\软著文档\用户手册_生成版_截图索引.md

Dependencies:
  pip install python-docx lxml

Notes:
- Screenshots are taken from the *source* docx (word/media/*) and re-inserted.
- Image appearance order is based on word/document.xml drawing order.
"""

from __future__ import annotations

import argparse
import csv
import os
import re
import shutil
import sys
import zipfile
from dataclasses import dataclass
from datetime import datetime

try:
    from lxml import etree
except Exception as e:
    print("[FATAL] Missing dependency: lxml. Please run: pip install lxml", file=sys.stderr)
    raise

try:
    import docx
    from docx.shared import Inches
except Exception as e:
    print("[FATAL] Missing dependency: python-docx. Please run: pip install python-docx", file=sys.stderr)
    raise


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
}


@dataclass
class ImgOccur:
    index: int
    rId: str
    target: str  # e.g. word/media/image1.png
    file_name: str  # image1.png
    context_before: str
    context_after: str
    paragraph_text: str


def read_zip_text(z: zipfile.ZipFile, name: str) -> bytes:
    try:
        return z.read(name)
    except KeyError:
        raise RuntimeError(f"Missing part in docx: {name}")


def norm_text(s: str) -> str:
    s = re.sub(r"\s+", " ", s or "").strip()
    return s


def iter_paragraphs_with_text_and_blips(doc_xml: bytes):
    root = etree.fromstring(doc_xml)
    # Iterate paragraphs in document order.
    for p in root.xpath("//w:body/w:p", namespaces=NS):
        # paragraph plain text
        texts = [t for t in p.xpath(".//w:t/text()", namespaces=NS)]
        p_text = norm_text("".join(texts))
        # blip embeds inside this paragraph
        blips = p.xpath(".//a:blip", namespaces=NS)
        embeds = []
        for b in blips:
            rid = b.get(f"{{{NS['r']}}}embed")
            if rid:
                embeds.append(rid)
        yield p_text, embeds


def parse_rels(rels_xml: bytes) -> dict[str, str]:
    root = etree.fromstring(rels_xml)
    mapping: dict[str, str] = {}
    for rel in root.xpath("//pr:Relationship", namespaces=NS):
        rid = rel.get("Id")
        target = rel.get("Target")
        if not rid or not target:
            continue
        # For images Target typically like: media/image1.png
        mapping[rid] = target
    return mapping


def extract_images_in_order(src_docx: str, out_media_dir: str) -> list[ImgOccur]:
    os.makedirs(out_media_dir, exist_ok=True)

    with zipfile.ZipFile(src_docx) as z:
        doc_xml = read_zip_text(z, "word/document.xml")
        rels_xml = read_zip_text(z, "word/_rels/document.xml.rels")
        rid_to_target = parse_rels(rels_xml)

        # Build paragraph list so we can provide context around each image.
        paragraphs: list[str] = []
        embeds_by_para: list[list[str]] = []
        for p_text, embeds in iter_paragraphs_with_text_and_blips(doc_xml):
            paragraphs.append(p_text)
            embeds_by_para.append(embeds)

        def find_prev_text(i: int) -> str:
            for j in range(i - 1, max(-1, i - 6), -1):
                if paragraphs[j]:
                    return paragraphs[j]
            return ""

        def find_next_text(i: int) -> str:
            for j in range(i + 1, min(len(paragraphs), i + 6)):
                if paragraphs[j]:
                    return paragraphs[j]
            return ""

        occ: list[ImgOccur] = []
        img_idx = 0
        for i, embeds in enumerate(embeds_by_para):
            if not embeds:
                continue
            for rid in embeds:
                target = rid_to_target.get(rid)
                if not target:
                    continue
                # Normalize target to word/media/xxx
                if target.startswith("../"):
                    target = target[3:]
                if not target.startswith("media/"):
                    # some other relationship
                    continue

                media_path = "word/" + target
                file_name = os.path.basename(target)

                # extract file to out_media_dir (keep original name)
                out_path = os.path.join(out_media_dir, file_name)
                if not os.path.exists(out_path):
                    with z.open(media_path) as src, open(out_path, "wb") as dst:
                        shutil.copyfileobj(src, dst)

                img_idx += 1
                occ.append(
                    ImgOccur(
                        index=img_idx,
                        rId=rid,
                        target=media_path,
                        file_name=file_name,
                        context_before=find_prev_text(i),
                        context_after=find_next_text(i),
                        paragraph_text=paragraphs[i],
                    )
                )
        return occ


def pick_image_for_module(images: list[ImgOccur], used: set[int], keywords: list[str]) -> int | None:
    """Pick a screenshot for a module.

    Priority:
    1) Unused screenshot whose nearby text matches any keyword
    2) Next unused screenshot by appearance order
    3) If all screenshots have been used, reuse the first screenshot (to satisfy "each module has a screenshot").
    """
    if not images:
        return None

    kw = [k for k in keywords if k]
    if kw:
        for img in images:
            if img.index in used:
                continue
            ctx = " ".join([img.context_before, img.paragraph_text, img.context_after])
            if any(k in ctx for k in kw):
                return img.index

    for img in images:
        if img.index not in used:
            return img.index

    # All used: allow reuse
    return images[0].index


def add_heading(document: docx.Document, text: str, level: int = 1):
    document.add_heading(text, level=level)


def add_bullets(document: docx.Document, items: list[str]):
    for it in items:
        p = document.add_paragraph(it, style="List Bullet")


def add_steps(document: docx.Document, steps: list[str]):
    for i, s in enumerate(steps, 1):
        document.add_paragraph(f"{i}. {s}", style="List Number")


def insert_image(document: docx.Document, media_dir: str, img: ImgOccur, caption: str):
    path = os.path.join(media_dir, img.file_name)
    if os.path.exists(path):
        document.add_picture(path, width=Inches(6.2))
        document.add_paragraph(f"图{img.index} {caption}")
        document.add_paragraph("（截图来源：原《用户手册.docx》）")
    else:
        document.add_paragraph(f"[缺失图片文件：{img.file_name}]")


def build_manual(root_dir: str, src_docx: str, out_docx: str):
    out_dir = os.path.dirname(out_docx)
    media_dir = os.path.join(out_dir, "_用户手册_media")
    if os.path.exists(media_dir):
        shutil.rmtree(media_dir)

    images = extract_images_in_order(src_docx, media_dir)

    # Define modules based on actual routes/pages found in Vue.
    modules = [
        {
            "title": "登录与退出（管理员/用户入口）",
            "keywords": ["登录", "管理员", "账号", "密码", "退出"],
            "content": {
                "desc": [
                    "系统分为患者入口（居民体质辨识）与管理员后台两类入口。",
                    "患者入口以“登记信息→四诊采集→生成报告”为主；管理员后台用于数据概览、患者管理、诊断记录查询等。",
                ],
                "steps": [
                    "打开浏览器，访问前端地址：http://localhost:5173/。",
                    "点击右上角「更多功能」→「管理后台」。",
                    "在管理员登录页输入用户名与密码，点击「登录」。",
                    "登录成功后进入管理后台；如需退出，点击侧边栏底部「退出」。",
                ],
                "notes": [
                    "若无法进入管理后台，通常是后端 8080 未启动或接口代理未通。",
                    "管理员登录状态保存在浏览器本机（localStorage），清理浏览器缓存会导致需要重新登录。",
                ],
                "faq": [
                    "提示‘网络错误，请检查后端服务’：请确认 demo 后端在 8080 端口运行。",
                    "点击后台页面自动跳回登录页：表示本机缺少 admin_token，重新登录即可。",
                ],
            },
        },
        {
            "title": "报告设置（机构信息/签发/免责声明/AI 提示词）",
            "keywords": ["报告设置", "机构名称", "签发医师", "免责声明", "提示词", "侧重点"],
            "content": {
                "desc": [
                    "在患者入口主页右上角「更多功能 ▾」中，可打开「报告设置」。",
                    "此处配置的机构信息会显示在每份四诊报告/健康体检报告的抬头处；同时支持选择 AI 报告生成的侧重点或自定义提示词模板。",
                    "所有设置默认保存在本机浏览器中（更换电脑/清理浏览器数据会丢失，需要重新设置）。",
                ],
                "steps": [
                    "在主页点击右上角「更多功能 ▾」→ 选择「报告设置」。",
                    "在「机构信息」中填写：机构名称（必填）、机构地址（选填）、联系电话（选填）。",
                    "在「报告信息」中填写：签发医师（选填）、有效期说明（选填）、免责声明（选填）。",
                    "在「AI提示词设置」中选择报告侧重点（不侧重/望诊/闻诊/问诊/切诊），或填写自定义提示词模板（选填）。",
                    "点击「保存设置」，看到提示“报告设置已保存”即完成。",
                ],
                "notes": [
                    "机构名称为必填项；不填写将无法保存。",
                    "自定义提示词建议由管理员维护，避免加入不当或与实际业务不符的内容。",
                ],
                "faq": [
                    "保存后无变化：请刷新报告页面重新生成/打开报告；旧报告不会自动回写抬头。",
                    "设置丢失：可能清理了浏览器缓存或更换了浏览器/电脑，需重新配置。",
                ],
            },
        },
        {
            "title": "居民信息登记（开始测试）",
            "keywords": ["确认个人信息", "身份证", "手动填写", "感应"],
            "content": {
                "desc": [
                    "在开始四诊采集前，需要先录入/确认居民基本信息。",
                    "支持两种录入方式：①手动填写；②身份证读卡感应（需要读卡服务）。",
                ],
                "steps": [
                    "在首页点击「开始测试」。",
                    "选择录入方式：手动填写 或 身份证感应。",
                    "按页面提示填写必填项（姓名、性别、出生日期、身份证号、住址）。",
                    "点击「下一步，进入四诊辨识」。",
                ],
                "notes": [
                    "身份证号需 18 位；信息将用于建立诊断会话与生成报告抬头。",
                    "如使用身份证感应，需确保读卡服务端口 9009 可用（详见附录）。",
                ],
                "faq": [
                    "读卡一直失败：检查读卡器连接/驱动，确认 IdCardReaderService 已启动。",
                    "无法下一步：请检查是否有必填项为空或身份证号长度不正确。",
                ],
            },
        },
        {
            "title": "四诊辨识中心（望/闻/问/切选择）",
            "keywords": ["四诊", "完成进度", "生成报告"],
            "content": {
                "desc": [
                    "四诊中心用于选择具体采集模块，并显示当前居民的完成进度。",
                    "完成任一项即可生成阶段性报告；完成四项后可生成更全面的综合报告。",
                ],
                "steps": [
                    "进入四诊中心后，点击任一模块卡片：望诊/闻诊/问诊/切诊。",
                    "系统会为本次居民创建诊断会话（caseId），并锁定当前就诊ID。",
                    "完成采集后返回四诊中心，查看进度条变化。",
                    "点击「生成报告」进入报告页。",
                ],
                "notes": [
                    "已完成的模块会被锁定，避免误操作覆盖结果。",
                    "若更换居民信息，请返回上一页重新登记。",
                ],
                "faq": [
                    "卡片提示已锁定：表示该居民已完成该项采集；需要重新测量请在模块内点击“重新分析/重新测量”。",
                ],
            },
        },
        {
            "title": "望诊（舌象采集与分析）",
            "keywords": ["望诊", "舌", "摄像头", "舌象"],
            "content": {
                "desc": [
                    "望诊模块通过摄像头拍照或上传图片采集舌象，调用 AI 引擎生成舌质/舌苔等分析结果。",
                ],
                "steps": [
                    "在四诊中心点击「望诊分析」。",
                    "选择「开启摄像头拍照」或「从相册选择图片」。",
                    "按取景框提示对准舌面，拍照/选择后等待 AI 分析完成。",
                    "阅读结果（量化指标、初步结论、模型解读），点击「确认并返回」。",
                ],
                "notes": [
                    "拍摄环境建议光线均匀、避免强反光；舌面尽量平伸。",
                    "如摄像头权限被禁用，请在浏览器站点设置中允许摄像头访问。",
                ],
                "faq": [
                    "一直转圈分析：检查 AI 服务是否运行（5000 端口）。",
                    "无法打开摄像头：检查浏览器权限/摄像头是否被其他软件占用。",
                ],
            },
        },
        {
            "title": "闻诊（音频录制与分析）",
            "keywords": ["闻诊", "录制", "音频", "麦克风"],
            "content": {
                "desc": [
                    "闻诊模块通过录制语音与呼吸声，提取音频特征并给出体质倾向与置信度。",
                ],
                "steps": [
                    "在四诊中心点击「闻诊分析」。",
                    "点击「开始录制」，先说出姓名，再进行 3~5 次深呼吸（约 15 秒）。",
                    "点击「停止录制」，可先播放确认清晰度。",
                    "点击「分析 & 提交」，等待结果生成并确认返回。",
                ],
                "notes": [
                    "请在安静环境录制，避免背景噪声（人群/风扇/音乐）。",
                    "若无麦克风或权限被禁用，请在浏览器设置中开启麦克风权限。",
                ],
                "faq": [
                    "没有波形/无法录音：检查麦克风权限；尝试更换浏览器。",
                ],
            },
        },
        {
            "title": "问诊（智能问卷采集）",
            "keywords": ["问诊", "问卷", "模板选择", "33"],
            "content": {
                "desc": [
                    "问诊模块提供标准体质问卷（33 题）与专项问诊模板，用于采集症状与生活习惯信息。",
                    "完成后系统自动生成问诊结论与养生建议，并可生成阶段性报告。",
                ],
                "steps": [
                    "在四诊中心点击「问诊分析」。",
                    "先在“模板选择”中选择问诊人群/场景（标准或专项模板）。",
                    "逐题选择答案（部分题目需要填写身高体重用于 BMI 计算）。",
                    "最后点击「提交并完成采集」，查看问诊结论与建议。",
                    "点击「确认并返回」回到四诊中心。",
                ],
                "notes": [
                    "答题过程中请尽量如实选择；专项模板题目较少，适合快速筛查。",
                ],
                "faq": [
                    "无法进入下一题：需要先选择本题答案。",
                ],
            },
        },
        {
            "title": "切诊（脉搏采集与算法分析）",
            "keywords": ["切诊", "脉诊", "PPG", "采集"],
            "content": {
                "desc": [
                    "切诊模块对接脉搏采集设备，采集 PPG 脉搏波形并计算心率、血氧、信号质量等指标，生成中医脉象建议。",
                ],
                "steps": [
                    "在四诊中心点击「切诊分析」。",
                    "引导居民将手指平稳放在传感器上，保持静止。",
                    "点击「开始切诊」，等待倒计时采集（默认 60 秒）。",
                    "采集完成后系统自动分析并生成报告，点击「确认并返回」。",
                ],
                "notes": [
                    "采集期间请保持手指静止；信号质量低会影响结论稳定性。",
                    "该模块依赖脉搏算法服务（8000 端口）与设备连接（详见硬件指引）。",
                ],
                "faq": [
                    "提示‘未检测到手指/干扰’：检查手指贴合、传感器清洁、环境光干扰。",
                ],
            },
        },
        {
            "title": "四诊合参诊断报告（导出/打印）",
            "keywords": ["诊断报告", "导出PDF", "打印", "四诊合参"],
            "content": {
                "desc": [
                    "报告模块融合望闻问切结果，生成综合分析文本，并支持导出 PDF 与打印。",
                    "当四诊未全部完成时，报告会显示完成度百分比与缺失提示。",
                ],
                "steps": [
                    "在四诊中心点击「生成报告」。",
                    "等待 AI 综合分析完成（支持流式预览）。",
                    "检查报告内容（患者信息、各诊断模块结果、综合建议）。",
                    "点击「导出PDF」保存文件，或点击「打印」输出纸质版。",
                ],
                "notes": [
                    "导出 PDF 依赖浏览器下载权限；如被拦截，请允许弹窗/下载。",
                    "报告抬头的机构信息可在首页「报告设置」中配置。",
                ],
                "faq": [
                    "导出 PDF 空白：尝试等待页面完全渲染后再导出；或更换 Chrome 浏览器。",
                ],
            },
        },
        {
            "title": "辨识档案管理（查询与详情）",
            "keywords": ["辨识档案", "查看详情", "只读"],
            "content": {
                "desc": [
                    "辨识档案用于集中查看所有居民的四诊辨识记录，支持按姓名/身份证搜索，并可打开抽屉查看详细结果。",
                    "该页面标记为“只读”，不支持编辑或删除历史记录。",
                ],
                "steps": [
                    "在首页顶部导航中切换到「辨识档案管理」。",
                    "在搜索框输入姓名或身份证号，点击「搜索」。",
                    "在列表中点击「查看详情」，在右侧抽屉查看望/闻/问/切数据。",
                ],
                "notes": [
                    "身份证号在列表中会做脱敏显示，保护隐私。",
                ],
                "faq": [
                    "列表为空：通常是系统尚无历史数据，或后端数据库未初始化。",
                ],
            },
        },
        {
            "title": "居民体检管理与健康体检报告",
            "keywords": ["体检", "健康体检报告", "新增体检"],
            "content": {
                "desc": [
                    "体检管理用于录入居民综合体检数据（血压、血糖、BMI 等），并可生成健康体检报告。",
                ],
                "steps": [
                    "在首页顶部导航中切换到「居民体检管理」。",
                    "点击「新增体检」录入基本信息与体检数据，保存档案。",
                    "在列表中点击「查看报告」，进入健康体检报告页。",
                    "在报告页可导出 PDF 或打印。",
                ],
                "notes": [
                    "标记 * 的字段为必填；其余选填项越完整，报告内容越丰富。",
                ],
                "faq": [
                    "保存失败：检查后端服务与数据库连接；确认必填项已填写。",
                ],
            },
        },
        {
            "title": "体质统计分析（数据概览）",
            "keywords": ["体质统计", "分布", "刷新"],
            "content": {
                "desc": [
                    "体质统计分析对已完成的辨识数据进行汇总，以图表形式展示体质分布与服务人次。",
                ],
                "steps": [
                    "在首页顶部导航中切换到「体质统计分析」。",
                    "等待数据加载完成后查看统计卡片与分布条形图。",
                    "如需更新，点击「刷新」。",
                ],
                "notes": [
                    "统计数据来源于后端 /api/admin/* 接口，需管理员/后台服务支持。",
                ],
                "faq": [
                    "一直显示 0：可能后端接口未返回数据或数据库中无记录。",
                ],
            },
        },
        {
            "title": "管理后台（数据概览/患者管理/诊断记录）",
            "keywords": ["管理后台", "患者管理", "诊断记录", "数据概览"],
            "content": {
                "desc": [
                    "管理后台用于系统运维与数据管理，包含：数据概览、患者管理（可删除患者）、诊断记录查看等。",
                ],
                "steps": [
                    "进入管理后台后，左侧选择「数据概览」查看统计数据。",
                    "选择「患者管理」可搜索患者并执行删除（谨慎）。",
                    "选择「诊断记录」可查看每次诊断的完成情况并打开详情抽屉。",
                ],
                "notes": [
                    "删除患者会同时删除其所有诊断记录，请谨慎操作。",
                ],
                "faq": [
                    "删除无反应：检查网络请求是否被拦截，或后端权限校验失败。",
                ],
            },
        },
        {
            "title": "系统介绍 / 硬件指引 / 中医文化",
            "keywords": ["系统介绍", "硬件", "中医文化"],
            "content": {
                "desc": [
                    "系统提供科普与引导页面：系统介绍、硬件指引、中医文化模块（图文阅读）。",
                ],
                "steps": [
                    "在首页点击右上角「更多功能」，进入对应页面。",
                    "在中医文化模块中选择专题阅读，支持查看图文内容。",
                ],
                "notes": [
                    "该模块用于展示内容与引导，不影响诊断主流程。",
                ],
                "faq": [
                    "页面图片不显示：检查前端静态资源是否部署完整，或后端 /uploads 代理是否可用。",
                ],
            },
        },
    ]

    document = docx.Document()

    # Title
    document.add_heading("中医体质辨识系统 用户手册（生成版）", 0)
    document.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph("说明：本手册根据原《用户手册.docx》截图顺序与项目实际功能自动生成。")

    add_heading(document, "1. 产品概述", 1)
    document.add_paragraph(
        "中医体质辨识系统面向基层健康管理与居民体质评估场景，"
        "通过‘望、闻、问、切’四诊信息采集与 AI 辅助分析，"
        "生成阶段性或四诊合参的体质辨识报告，并支持档案管理、体检管理与统计分析。"
    )

    add_heading(document, "2. 运行环境", 1)
    add_bullets(document, [
        "操作系统：Windows 10/11（建议 64 位）",
        "浏览器：Chrome（推荐）/Edge（Chromium 内核）",
        "网络：本机部署可离线运行；如跨机部署需保证端口可达",
        "依赖组件：Java 21 + Maven、Python 3.10+、Node.js 16+（用于前端）",
        "硬件（可选）：身份证读卡器、脉搏采集传感器、摄像头、麦克风",
    ])

    add_heading(document, "3. 安装与启动", 1)
    document.add_paragraph("本项目提供一键启动脚本，推荐按以下方式启动：")
    add_steps(document, [
        "双击运行仓库根目录 start.vbs（推荐，自动释放端口并后台启动）。",
        "或运行 start_tcm.bat（会打开多个 cmd 窗口）。",
        "等待前端启动完成后浏览器访问：http://localhost:5173",
    ])

    add_heading(document, "3.1 服务组成、端口与访问地址", 2)
    add_bullets(document, [
        "前端 Vue3/Vite：http://localhost:5173（页面入口）",
        "后端 Spring Boot（demo）：http://localhost:8080（统一 /api 接口）",
        "AI 服务（tcm-ai-service，Python）：http://localhost:5000（AI 推理/分析）",
        "脉搏算法服务（pulse2，Python）：http://localhost:8000（PPG/脉诊算法）",
        "身份证读卡服务（IdCardReaderService）：http://localhost:9009（读卡服务）",
    ])

    document.add_paragraph("注意：前端开发服务器会把 /api 请求代理到 8080，把 /uploads 代理到 8080。")

    add_heading(document, "4. 登录与退出", 1)
    document.add_paragraph("患者入口无需登录；管理后台需要管理员登录。具体操作见后续章节。")

    add_heading(document, "5. 主要功能模块说明", 1)

    used: set[int] = set()
    screenshot_index_rows = []

    for m in modules:
        add_heading(document, m["title"], 2)

        img_idx = pick_image_for_module(images, used, m.get("keywords", []))
        if img_idx is not None:
            used.add(img_idx)
            img = next(i for i in images if i.index == img_idx)
            caption = m["title"]
            insert_image(document, media_dir, img, caption)
            screenshot_index_rows.append({
                "图号": f"图{img.index}",
                "图片文件": img.file_name,
                "章节": m["title"],
                "归类依据": f"上下文命中关键词：{','.join(m.get('keywords', [])[:6])}；近邻文本：{(img.context_before or img.paragraph_text or img.context_after)[:80]}"
            })
        else:
            document.add_paragraph("（未找到可用截图，后续请补充）")

        for d in m["content"].get("desc", []):
            document.add_paragraph(d)

        if m["content"].get("steps"):
            document.add_paragraph("操作步骤：")
            add_steps(document, m["content"]["steps"])

        if m["content"].get("notes"):
            document.add_paragraph("注意事项：")
            add_bullets(document, m["content"]["notes"])

        if m["content"].get("faq"):
            document.add_paragraph("常见问题：")
            add_bullets(document, m["content"]["faq"])

    # Remaining screenshots
    remaining = [img for img in images if img.index not in used]
    add_heading(document, "6. 界面截图汇总（未能自动判别）", 1)
    if remaining:
        document.add_paragraph("以下截图来自原手册，但未能通过上下文关键字准确归类到具体功能模块，供人工核对：")
        for img in remaining:
            insert_image(document, media_dir, img, "界面截图（待人工确认）")
            screenshot_index_rows.append({
                "图号": f"图{img.index}",
                "图片文件": img.file_name,
                "章节": "界面截图汇总（未能自动判别）",
                "归类依据": f"未命中关键词；近邻文本：{(img.context_before or img.paragraph_text or img.context_after)[:80]}"
            })
    else:
        document.add_paragraph("所有截图已分配到功能章节。")

    add_heading(document, "7. 常见问题 / 故障排查", 1)
    add_bullets(document, [
        "前端打不开（5173）：检查是否已执行 npm install；端口是否被占用；是否启动了 npm run dev。",
        "接口报错（/api 失败）：确认 Spring Boot 后端 8080 正常运行；检查代理配置。",
        "AI 分析失败：确认 tcm-ai-service 5000 启动；检查 Python 环境与依赖是否安装。",
        "切诊无波形：确认 pulse2 8000 启动；设备连接是否正常；信号质量是否过低。",
        "身份证读卡失败：确认 IdCardReaderService 9009 启动；读卡器驱动与连接正常。",
    ])

    add_heading(document, "8. 附录：端口/服务清单", 1)
    add_bullets(document, [
        "8080：Spring Boot 后端（统一业务 API）",
        "5173：Vue3/Vite 前端（用户界面入口）",
        "5000：AI 服务（望/闻/问等 AI 推理）",
        "8000：脉搏算法服务（切诊/脉诊）",
        "9009：身份证读卡服务",
    ])

    document.add_paragraph("提示：可用 stop.vbs 尝试一键结束相关后台进程并释放端口。")

    # Save
    os.makedirs(os.path.dirname(out_docx), exist_ok=True)
    document.save(out_docx)

    # Write index tables
    csv_path = os.path.join(out_dir, "用户手册_生成版_截图索引.csv")
    md_path = os.path.join(out_dir, "用户手册_生成版_截图索引.md")

    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=["图号", "图片文件", "章节", "归类依据"])
        w.writeheader()
        for row in screenshot_index_rows:
            w.writerow(row)

    with open(md_path, "w", encoding="utf-8") as f:
        f.write("# 用户手册_生成版 截图索引\n\n")
        f.write("| 图号 | 图片文件 | 章节 | 归类依据 |\n")
        f.write("|---|---|---|---|\n")
        for row in screenshot_index_rows:
            f.write(f"| {row['图号']} | {row['图片文件']} | {row['章节']} | {row['归类依据'].replace('|','|')} |\n")

    print(f"[OK] Generated: {out_docx}")
    print(f"[OK] Screenshot index: {csv_path}")
    print(f"[OK] Screenshot index: {md_path}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", required=True, help="source docx with screenshots")
    ap.add_argument("--out", required=True, help="output docx")
    ap.add_argument("--root", default="", help="project root (reserved)")
    args = ap.parse_args()

    if not os.path.exists(args.src):
        raise SystemExit(f"Source docx not found: {args.src}")

    build_manual(args.root, args.src, args.out)


if __name__ == "__main__":
    main()
