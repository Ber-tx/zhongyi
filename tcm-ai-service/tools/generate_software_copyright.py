# -*- coding: utf-8 -*-
"""
软件著作权文档生成工具
- 生成前端源代码文档（30页×50行，无注释）
- 生成后端源代码文档（30页×50行，无注释）

使用方法:
    pip install python-docx
    python generate_software_copyright.py
"""

import os
import re
import io
import tokenize
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# ==================== 配置区域 ====================

# 项目根目录
PROJECT_ROOT = Path(r"E:\项目")

# 前端源码目录
FRONTEND_DIR = PROJECT_ROOT / "Vue" / "zhongyi" / "src"

# 后端源码目录 (Java Spring Boot)
BACKEND_DIR = PROJECT_ROOT / "demo" / "src" / "main" / "java" / "com" / "tx" / "demo"

# Python后端目录（混排到后端代码文档）
PYTHON_BACKEND_DIR = PROJECT_ROOT / "tcm-ai-service"

# 输出目录
OUTPUT_DIR = PROJECT_ROOT / "软著文档"

# 每页行数
LINES_PER_PAGE = 50

# 需要的总页数
TOTAL_PAGES = 30

# 总行数
TOTAL_LINES = LINES_PER_PAGE * TOTAL_PAGES  # 1500行

# 单行最大字符数（仅代码内容，不含左侧行号）
# 取更保守值，避免Word自动换行导致“行号不增加但内容换行”
MAX_CHARS_PER_LINE = 48

# 是否显示每页标题（关闭可避免占位和大段空白）
SHOW_PAGE_TITLE = False

# 软件名称
SOFTWARE_NAME = "中医智能诊断系统"

# ==================== 前端文件选择（按优先级排序） ====================

FRONTEND_FILES = [
    # 核心页面组件
    "views/Detect/DetectSelect.vue",
    "views/Detect/Wang.vue",
    "views/Detect/Wen.vue",
    "views/Detect/Wenjuan.vue",
    "views/Detect/Qie.vue",
    "views/Report.vue",
    "views/Home.vue",
    "views/HealthReport.vue",
    # TCM模块
    "views/TCM/PatientRegister.vue",
    "views/TCM/DiagnosisArchive.vue",
    "views/TCM/HealthExam.vue",
    "views/TCM/ConstitutionStats.vue",
    # 管理员模块
    "views/Admin/AdminDashboard.vue",
    "views/Admin/AdminLogin.vue",
    # 其他页面
    "views/SystemIntro.vue",
    "views/HardwareGuide.vue",
    # 文化模块
    "views/Culture/CultureHome.vue",
    "views/Culture/ModuleFour.vue",
    "views/Culture/ModuleTen.vue",
    # API接口
    "api/detect.js",
    "api/report.js",
    "api/auth.js",
    "api/admin.js",
    "api/hardware.js",
    # 路由和工具
    "router/index.js",
    "utils/reportUtils.js",
    "constants/questionnaireTemplates.js",
    "constants/algorithmReferences.js",
    # 入口文件
    "main.js",
    "App.vue",
]

# ==================== 后端文件选择（按优先级排序） ====================

BACKEND_FILES = [
    # Controller层
    "controller/WangController.java",
    "controller/QuestionController.java",
    "controller/AdminController.java",
    "controller/PatientController.java",
    "controller/DiagnosisSessionController.java",
    "controller/HealthExamController.java",
    "controller/IdCardDemoController.java",
    # Service层
    "service/impl/QuestionServiceImpl.java",
    "service/impl/PatientServiceImpl.java",
    "service/QuestionService.java",
    "service/PatientService.java",
    "service/HealthExamService.java",
    # Mapper层
    "mapper/DiagnosisMapper.java",
    "mapper/AdminMapper.java",
    "mapper/HealthExamMapper.java",
    "mapper/PatientMapper.java",
    "mapper/QuestionMapper.java",
    # Entity层
    "entity/Diagnosis.java",
    "entity/HealthExam.java",
    "entity/Patient.java",
    "entity/Question.java",
    "entity/Answer.java",
    "entity/Admin.java",
    # 配置和工具
    "config/WebConfig.java",
    "config/JacksonConfig.java",
    "config/HttpClientConfig.java",
    "utils/Result.java",
    "utils/JwtUtils.java",
    "interceptor/AdminInterceptor.java",
    "vo/IdCardInfoVO.java",
    # 入口
    "ZhongyiApplication.java",
]

PYTHON_BACKEND_FILES = [
    "main.py",
    "api/router_synthesis.py",
    "api/router_tongue.py",
    "core/tongue_shizhen.py",
    "core/image_processor.py",
    "core/visualizer.py",
]

# ==================== 注释删除函数 ====================

def remove_vue_js_comments(code: str) -> str:
    """删除Vue/JS代码中的注释"""
    code = re.sub(r'<!--[\s\S]*?-->', '', code)
    code = re.sub(r'/\*[\s\S]*?\*/', '', code)
    lines = code.split('\n')
    result_lines = []
    in_block_comment = False

    for line in lines:
        stripped = line.strip()

        # 处理块注释
        if in_block_comment:
            if '*/' in line:
                in_block_comment = False
                # 处理 */ 之后的内容
                idx = line.find('*/')
                if idx + 2 < len(line):
                    remaining = line[idx + 2:].strip()
                    if remaining and not remaining.startswith('//'):
                        result_lines.append(remaining)
            continue

        # 跳过单行注释
        if stripped.startswith('//'):
            continue

        # 检测块注释开始
        if stripped.startswith('/*'):
            if '*/' in line:
                # 单行块注释，跳过整行
                continue
            else:
                in_block_comment = True
                continue

        # 删除行内注释
        if '//' in line:
            # 简单处理：假设 // 后面都是注释
            idx = line.find('//')
            # 确保不是字符串中的 //
            before = line[:idx]
            if before.count('"') % 2 == 0 and before.count("'") % 2 == 0:
                line = line[:idx].rstrip()

        # 跳过空行（可选，如果需要保留空行可以注释掉）
        # if not stripped:
        #     continue

        result_lines.append(line)

    return '\n'.join(result_lines)


def remove_java_comments(code: str) -> str:
    """删除Java代码中的注释"""
    code = re.sub(r'/\*[\s\S]*?\*/', '', code)
    lines = code.split('\n')
    result_lines = []
    in_block_comment = False

    for line in lines:
        stripped = line.strip()

        # 处理块注释
        if in_block_comment:
            if '*/' in line:
                in_block_comment = False
                idx = line.find('*/')
                if idx + 2 < len(line):
                    remaining = line[idx + 2:].strip()
                    if remaining and not remaining.startswith('//'):
                        result_lines.append(remaining)
            continue

        # 跳过单行注释
        if stripped.startswith('//'):
            continue

        # 跳过Javadoc注释
        if stripped.startswith('/**') or stripped.startswith('*'):
            if stripped.startswith('/**'):
                if '*/' in line:
                    continue
                else:
                    in_block_comment = True
            continue

        # 检测块注释开始
        if stripped.startswith('/*'):
            if '*/' in line:
                continue
            else:
                in_block_comment = True
                continue

        # 删除行内注释
        if '//' in line:
            idx = line.find('//')
            before = line[:idx]
            if before.count('"') % 2 == 0:
                line = line[:idx].rstrip()

        result_lines.append(line)

    return '\n'.join(result_lines)


def remove_python_comments(code: str) -> str:
    """删除Python代码中的注释"""
    out = []
    prev_toktype = tokenize.INDENT
    try:
        tokens = tokenize.generate_tokens(io.StringIO(code).readline)
        for tok_type, tok_str, _, _, _ in tokens:
            if tok_type == tokenize.COMMENT:
                continue
            if tok_type == tokenize.STRING and prev_toktype in (
                tokenize.INDENT, tokenize.NEWLINE, tokenize.DEDENT, tokenize.ENCODING
            ):
                continue
            out.append((tok_type, tok_str))
            prev_toktype = tok_type
    except Exception:
        lines = code.split('\n')
        result_lines = []
        for line in lines:
            stripped = line.strip()
            if not stripped or stripped.startswith('#'):
                continue
            if '#' in line:
                idx = line.find('#')
                before = line[:idx]
                if before.count('"') % 2 == 0 and before.count("'") % 2 == 0:
                    line = line[:idx].rstrip()
            result_lines.append(line)
        return '\n'.join(result_lines)
    return tokenize.untokenize(out)


# ==================== 代码提取函数 ====================

def read_file_content(file_path: Path) -> str:
    """读取文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"读取文件失败: {file_path}, 错误: {e}")
        return ""


def extract_frontend_code() -> str:
    """提取前端代码"""
    all_code = []

    for rel_path in FRONTEND_FILES:
        file_path = FRONTEND_DIR / rel_path
        if file_path.exists():
            print(f"读取前端文件: {rel_path}")
            code = read_file_content(file_path)
            if code:
                # 删除注释
                clean_code = remove_vue_js_comments(code)
                all_code.append(clean_code)
        else:
            print(f"文件不存在: {rel_path}")

    return '\n'.join(all_code)


def extract_backend_code() -> str:
    """提取后端代码"""
    all_code = []

    for rel_path in BACKEND_FILES:
        file_path = BACKEND_DIR / rel_path
        if file_path.exists():
            print(f"读取后端文件: {rel_path}")
            code = read_file_content(file_path)
            if code:
                # 删除注释
                clean_code = remove_java_comments(code)
                all_code.append(clean_code)
        else:
            print(f"文件不存在: {rel_path}")

    for rel_path in PYTHON_BACKEND_FILES:
        file_path = PYTHON_BACKEND_DIR / rel_path
        if file_path.exists():
            print(f"读取Python后端文件: {rel_path}")
            code = read_file_content(file_path)
            if code:
                clean_code = remove_python_comments(code)
                all_code.append(clean_code)
        else:
            print(f"文件不存在: {rel_path}")

    return '\n'.join(all_code)


def format_code_to_lines(code: str, target_lines: int = TOTAL_LINES) -> list:
    """格式化代码为指定行数，删除注释残留与空行，并按可见行切分长代码"""
    lines = code.split('\n')

    def fast_wrap_line(line: str, width: int) -> list:
        if len(line) <= width:
            return [line]
        indent = re.match(r'^\s*', line).group(0)
        cont_prefix = indent + "  "
        next_width = max(8, width - len(cont_prefix))
        chunks = []
        chunks.append(line[:width])
        rest = line[width:]
        while rest:
            part = rest[:next_width]
            chunks.append(cont_prefix + part)
            rest = rest[next_width:]
        return chunks

    cleaned_lines = []
    for raw in lines:
        if len(cleaned_lines) >= target_lines:
            break
        line = raw.expandtabs(4).rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("//") or stripped.startswith("#"):
            continue
        if stripped.startswith("/*") or stripped.startswith("*/") or stripped.startswith("*"):
            continue
        # 先按可见宽度拆分，确保Word不会自动折行
        wrapped = fast_wrap_line(line, MAX_CHARS_PER_LINE)
        if wrapped:
            for seg in wrapped:
                seg = seg.rstrip()
                if seg:
                    cleaned_lines.append(seg)
                    if len(cleaned_lines) >= target_lines:
                        break

    print(f"有效代码行数: {len(cleaned_lines)}")

    if not cleaned_lines:
        cleaned_lines = ["pass"]

    if len(cleaned_lines) < target_lines:
        print(f"代码行数不足，自动补齐到 {target_lines} 行")
        idx = 0
        while len(cleaned_lines) < target_lines:
            cleaned_lines.append(cleaned_lines[idx % len(cleaned_lines)])
            idx += 1
    return cleaned_lines[:target_lines]


# ==================== 文档生成函数 ====================

def create_source_code_document(title: str, code_lines: list, output_path: Path):
    """创建源代码文档"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)  # 五号字体

    # 强制严格页数与每页行数
    if len(code_lines) != TOTAL_LINES:
        raise ValueError(f"{title} 代码行数必须为 {TOTAL_LINES}，当前为 {len(code_lines)}")
    total_pages = TOTAL_PAGES

    # 按页生成内容
    for page_num in range(total_pages):
        start_idx = page_num * LINES_PER_PAGE
        end_idx = start_idx + LINES_PER_PAGE
        page_lines = code_lines[start_idx:end_idx]

        if SHOW_PAGE_TITLE:
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header.add_run(f"{SOFTWARE_NAME} - {title}")
            run.font.size = Pt(12)
            run.font.bold = True

        # 添加代码内容
        for i, line in enumerate(page_lines):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = Pt(12)
            para.paragraph_format.keep_together = False
            para.paragraph_format.keep_with_next = False
            para.paragraph_format.widow_control = False
            # 添加代码
            code_run = para.add_run(line)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(9.5)

        # 分页（最后一页不分页）
        if page_num < total_pages - 1:
            doc.add_page_break()

    # 保存文档
    doc.save(output_path)
    print(f"文档已保存: {output_path}")
    print(f"总页数: {total_pages}, 总行数: {len(code_lines)}")


# ==================== 主函数 ====================

def main():
    print("=" * 60)
    print(f"{SOFTWARE_NAME} - 软著文档生成工具")
    print("=" * 60)

    # 创建输出目录
    OUTPUT_DIR.mkdir(exist_ok=True)

    # 1. 生成前端源代码文档
    print("\n[1/2] 生成前端源代码文档...")
    frontend_code = extract_frontend_code()
    frontend_lines = format_code_to_lines(frontend_code)
    create_source_code_document(
        "前端源代码",
        frontend_lines,
        OUTPUT_DIR / f"{SOFTWARE_NAME}_前端源代码.docx"
    )

    # 2. 生成后端源代码文档
    print("\n[2/2] 生成后端源代码文档...")
    backend_code = extract_backend_code()
    backend_lines = format_code_to_lines(backend_code)
    create_source_code_document(
        "后端源代码",
        backend_lines,
        OUTPUT_DIR / f"{SOFTWARE_NAME}_后端源代码.docx"
    )

    print("\n" + "=" * 60)
    print("文档生成完成！")
    print(f"输出目录: {OUTPUT_DIR}")
    print("=" * 60)


if __name__ == "__main__":
    main()
